'''import lancedb
from docx import Document
import os

# --- Configuration ---
DB_PATH = "data/lancedb"         # Replace with your LanceDB path
TABLE_NAME = "testing_table"
TEXT_COLUMN = "text"             # Column that contains the text
OUTPUT_DOCX = "output_texts.docx"

# --- Connect to LanceDB ---
db = lancedb.connect(DB_PATH)
table = db.open_table(TABLE_NAME)

# --- Create Word Document ---
doc = Document()
doc.add_heading('Exported Texts from LanceDB', 0)

# --- Fetch and write texts ---
rows = table.to_pandas()

for i, text in enumerate(rows[TEXT_COLUMN]):
    try:
        clean_text = str(text).encode("utf-8", "ignore").decode("utf-8")  # Handle encoding
        doc.add_paragraph(f"{i+1}. {clean_text}")
    except Exception as e:
        print(f"Error on row {i}: {e}")

# --- Save the Document ---
doc.save(OUTPUT_DOCX)
print(f"Saved to {OUTPUT_DOCX}")
'''

import lancedb.embeddings as ld
import os
os.environ["JINA_API_KEY"]="jina_c4b2a33bd5794ddd98cf5f3905ce632bvPvHYr2uy3e3Eq9XUjBKR-yKV78e"

p=ld.JinaEmbeddings()
p=p.generate_text_embeddings("hi hello")
print(type(p[0]))
print(p[0])

'''import json
chunks=[{"text":"edfd","vector":234}]
f = open("data/vector.json", "w")
json.dump(chunks,f, indent=2)
f.close()
'''

